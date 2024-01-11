# Import necessary libraries
import docx
from morse_audio_decoder.morse import MorseCode

# Define Morse Code dictionary
MORSE_CODE_DICT = {'A':'.-', 'B':'-...',
                    'C':'-.-.', 'D':'-..', 'E':'.',
                    'F':'..-.', 'G':'--.', 'H':'....',
                    'I':'..', 'J':'.---', 'K':'-.-',
                    'L':'.-..', 'M':'--', 'N':'-.',
                    'O':'---', 'P':'.--.', 'Q':'--.-',
                    'R':'.-.', 'S':'...', 'T':'-',
                    'U':'..-', 'V':'...-', 'W':'.--',
                    'X':'-..-', 'Y':'-.--', 'Z':'--..',
                    '1':'.----', '2':'..---', '3':'...--',
                    '4':'....-', '5':'.....', '6':'-....',
                    '7':'--...', '8':'---..', '9':'----.',
                    '0':'-----', ', ':'--..--', '.':'.-.-.-',
                    '?':'..--..', '/':'-..-.', '-':'-....-',
                    '(':'-.--.', ')':'-.--.-', ' ':'       ',
                    "'":'.----.'}  # Morse code for apostrophe

# Function to convert Morse code to text
def morse_to_text(morse_code):
    # Replace Morse code for apostrophe with a unique placeholder
    morse_code = morse_code.replace('.----.', ' APOSTROPHE ')
    # Split the Morse code into words
    words = morse_code.split(" / ")
    decoded_words = []
    for word in words:
        # Split each word into letters
        letters = word.split(" ")
        # Decode each letter and join them to form a word
        decoded_word = "".join(MORSE_CODE_DICT.get(letter, '') if letter != 'APOSTROPHE' else "'" for letter in letters)
        decoded_words.append(decoded_word)
    # Join the decoded words with spaces to form a sentence
    return " ".join(decoded_words)

# Function to process audio file and convert to Morse code
def audio_to_morse(audio_file):
    # Create a MorseCode object from the audio file
    morse_code = MorseCode.from_wavfile(audio_file)
    # Decode the Morse code into text
    return morse_code.decode()

# Function to save text to Word document
def save_to_word(morse_code, text, filename):
    # Create a new Word document
    doc = docx.Document()
    # Add the Morse code and its English translation to the document
    doc.add_paragraph("Morse Code:")
    doc.add_paragraph(morse_code)
    doc.add_paragraph("Decoded Text:")
    doc.add_paragraph(text)
    # Save the document
    doc.save(filename)

# Main function
def main():
    # Prompt the user to enter the path to their .wav audio file
    audio_file = input("Please enter the path to your .wav audio file: ")
    # Convert the audio file to Morse code
    morse_code = audio_to_morse(audio_file)
    print(f"Decoded Morse Code: {morse_code}")  # Print the decoded Morse code
    # Convert the Morse code to text
    text = morse_to_text(morse_code)
    print(f"Decoded Text: {text}")
    # Prompt the user to enter the path where they want to save the Morse code and its English translation as a .docx Word document
    filename = input("Please enter the path where you want to save the Morse code and its English translation as a .docx Word document: ")
    save_to_word(morse_code, text, filename)

if __name__ == "__main__":
    main()
